#!/usr/bin/env python3
"""
AI Word Editor with DOCX, LaTeX, and PDF outputs.

Input:
    One Microsoft Word .docx file.

Outputs:
    - revised Brazilian Portuguese Word file
    - revised American English Word file
    - formatted Brazilian Portuguese PDF
    - formatted American English PDF
    - LaTeX source files and extracted figures/tables used to build the PDFs

The script preserves detected Word tables and embedded figures by replacing them
with stable placeholders during AI revision, then reinserting them into DOCX and
PDF outputs.

Fixes applied vs. v1:
    - latex_escape: backslash handled via placeholder to prevent double-escaping
    - is_heading_line: stricter heuristic (80 chars, no multi-sentence, no lowercase start)
    - text_to_latex_body: numbered-heading detection for multi-level hierarchy
    - call_claude: accepts optional system_prompt parameter
    - build_visual_audit: passes VISUAL_AUDIT_PROMPT as system, not user message
    - add_table_to_docx: three-line academic style (no vertical borders)
    - write_reportlab_pdf: three-line style via LINEABOVE/LINEBELOW
    - table_to_latex: content-aware column widths instead of uniform colspec
    - harmonize: chunked for large documents to avoid context-window overflow
    - requirements.txt: version ranges with upper bounds
"""

from __future__ import annotations

import argparse
import base64
import html
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Literal

from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image as RLImage,
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
except ImportError:
    colors = None

load_dotenv()

CONFIG = {
    "model": os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
    "max_chunk_chars": int(os.getenv("MAX_CHUNK_CHARS", "22000")),
    "max_output_tokens": int(os.getenv("MAX_OUTPUT_TOKENS", "20000")),
    "temperature": float(os.getenv("TEMPERATURE", "0.2")),
    "output_dir": Path(os.getenv("OUTPUT_DIR", "output")),
}

AREA_REFERENCES = """
Use approximate comparisons only when they help readers understand hectares:
- Brazil: 851.6 million ha
- Brazilian Amazon biome: about 419 million ha
- Brazilian Legal Amazon: about 502 million ha
- Mato Grosso state: about 90.3 million ha
- Para state: about 124.6 million ha
- Sao Paulo state: about 24.8 million ha
- France: about 55.1 million ha
- Germany: about 35.7 million ha
- United Kingdom: about 24.4 million ha
- Portugal: about 9.2 million ha
- Denmark: about 4.3 million ha
- Netherlands: about 4.2 million ha
- Belgium: about 3.1 million ha
"""

SYSTEM_PROMPT = f"""
You are an elite bilingual editor for academic and policy writing.

Revise the text for a mixed audience: informed non-specialists, policy readers,
journalists, practitioners, and specialists who still expect technical accuracy.

Use these references as editorial guidance, but do not mention them in the text:

- William Zinsser, On Writing Well: clarity, economy, and cutting clutter.
- Joseph Williams and Joseph Bizup, Style: Lessons in Clarity and Grace: clear subjects, strong verbs, known-to-new information flow, and graceful cohesion.
- Helen Sword, Stylish Academic Writing: rigorous academic prose without bureaucratic stiffness.
- Gerald Graff and Cathy Birkenstein, They Say / I Say: clear positioning, contribution, and argumentative movement.
- MEAL Plan: Main idea, Evidence, Analysis, Link in every paragraph where possible.
- Randy Olson, Houston, We Have a Narrative: context, problem, consequence, and resolution.
- Nancy Baron, Escape from the Ivory Tower: science and policy communication for public audiences.
- Roy Peter Clark, Writing Tools: rhythm, concrete language, strong openings, and memorable endings.
- The Economist style principles: precision, concision, confidence without hype.
- The New York Times explanatory journalism style: context, narrative flow, and clear digestion of evidence.
- Chip Heath and Dan Heath, Made to Stick: simplicity, concreteness, credibility, and memorable framing.
- Edward Tufte, Alberto Cairo, Cole Nussbaumer Knaflic, and Jonathan Schwabish: truthful and accessible communication of numbers and data.

Editorial rules:

1. Preserve every factual claim, number, institution, location, date, causal statement, and limitation unless the original is clearly ambiguous.
2. Do not invent evidence, references, citations, results, or policy claims.
3. Improve the text, not the argument. Keep the author's intellectual voice.
4. Make large numbers easier to understand. Keep the original number and add intuitive area comparisons only when useful.
5. Translate or explain jargon when a lay but informed reader would likely need help.
6. Use a natural academic tone: precise, direct, and readable.
7. Avoid robotic prose: no generic openings, no stock transitions, no inflated adjectives without evidence, no repetitive paragraph rhythm, no empty conclusions, no over-polished AI-like phrasing.
8. Vary sentence length and paragraph structure naturally.
9. Keep nuance, uncertainty, and caveats when they exist in the original.
10. Do not use em dashes.
11. Do not use excessive bullet points.
12. Do not mention AI, prompts, models, or the revision process.
13. Return only the revised text.
14. Keep all placeholders exactly unchanged. Placeholders look like [[TABLE:table_001]] and [[FIGURE:figure_001]]. Do not translate, remove, rename, or reorder them unless the surrounding text clearly requires moving them with the same paragraph.

Area comparison reference:
{AREA_REFERENCES}
"""

VISUAL_AUDIT_PROMPT = """
You are a senior scientific editor and data visualization specialist reviewing a Word document produced by researchers at a Brazilian environmental institution. The document targets academic peers, policy makers, journalists, development bank analysts, and engaged general public readers.

Your task is to audit every detected figure, table, chart, and map. Be concrete and direct. Do not provide generic advice. If you cannot directly inspect a visual element, state the limitation and still provide the strongest possible recommendations based on the caption, surrounding text, table values, and document context.

For each visual, use this exact structure:

---

VISUAL [number]: [type] -- [title or reference as it appears in text]

DIAGNOSIS (2-4 sentences):
State what the visual currently communicates, what it fails to communicate, and why. Be direct.

CHART/MAP TYPE:
State whether the chosen visual type is appropriate. If not, name the better alternative and explain in one sentence why it encodes the data more honestly or efficiently.

DATA INTEGRITY:
Flag any of the following if present or likely: truncated axes that exaggerate differences, dual y-axes that imply false correlation, area/volume encoding of linear quantities, choropleth of unnormalized absolute counts, rainbow or jet colormaps that introduce false gradients, 3D effects on 2D data, pie charts with more than four slices, or any encoding that would fail a data-ink efficiency test.

TITLE AND LABELS:
The title must make an assertion, not describe axes. The subtitle must orient the reader with units, geography, and time period. Source line must be in smaller type at bottom, formatted as: Source: [Institution], [Year]. Processed by [authors].

COLOR:
State whether the palette appears perceptually uniform, accessible to red-green colorblind readers, and interpretable in grayscale print. Recommend a specific named palette from ColorBrewer, Okabe-Ito, or Viridis. For deforestation and land use maps, recommend a sequential or diverging ColorBrewer palette and specify the class-break logic: equal interval, quantile, or Jenks natural breaks.

ANNOTATIONS:
Figures must be self-contained. Required elements: declarative title, subtitle with context, direct labels where possible, at least one annotation calling out the most important pattern, and source or processing notes.

CAPTION:
Write a revised caption of 2-4 sentences. Sentence 1 says what the visual shows. Sentence 2 gives the key finding. Sentence 3, if needed, gives a methodological note or caveat. Sentence 4, if needed, links the visual to the document's argument.

MAP-SPECIFIC RULES (only if the visual is a map):
- State the projection used or recommend one. For area comparisons in Brazil, recommend SIRGAS 2000 / Albers Equal Area (ESRI:102033). For global context, recommend Robinson or Winkel Tripel. Never recommend Mercator for thematic maps.
- The map must include a scale bar, north arrow, legend with explicit class breaks, and a locator inset.
- Choropleth maps must normalize by area or population unless the argument specifically requires totals.
- Use four to six classes maximum and state the classification method.
- Use a minimal basemap, thin gray administrative boundaries, and only analytically relevant rivers or roads.

TABLE-SPECIFIC RULES (only if the visual is a table):
- Use three horizontal lines only: above header row, below header row, below last data row.
- Use no vertical lines and no internal horizontal lines.
- Align numbers right and text left.
- State units once in the column header.
- If the table has more than eight rows, recommend converting it to a figure or splitting it.
- Highlight the key row or cell using bold or subtle shading.
- Write a table note below the table for methodological caveats, following APA 7 and three-line table convention.

CROSS-REFERENCE IN TEXT:
Identify the surrounding sentence that refers to this visual. Rewrite it so the text tells the reader what to conclude, while the visual provides the evidence.

PRIORITY:
Rate the revision urgency as HIGH, MEDIUM, or LOW.

---

After auditing all visuals, produce a VISUAL COHERENCE SUMMARY:
- Do all figures use the same color system? State inconsistencies.
- Are font sizes consistent across figures?
- Is the visual hierarchy consistent?
- List any figure that could be removed without loss of analytical content.
- Recommend the optimal figure count and sequence for the document's argument.

Do not mention the references that inform this audit. Apply them silently.
"""

LANGUAGE_NAME = {
    "pt-BR": "Brazilian Portuguese",
    "en-US": "American English",
}

# Styles Word uses for headings -- used when block.style is available
HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title", "Subtitle"}

BlockType = Literal["paragraph", "table", "figure"]


@dataclass
class FigureAsset:
    figure_id: str
    filename: str
    source_path: Path
    caption: str = ""


@dataclass
class TableAsset:
    table_id: str
    rows: list[list[str]]
    caption: str = ""


@dataclass
class Block:
    type: BlockType
    text: str = ""
    style: str = "Normal"
    ref_id: str = ""


@dataclass
class ExtractedDocument:
    title: str
    blocks: list[Block] = field(default_factory=list)
    figures: dict[str, FigureAsset] = field(default_factory=dict)
    tables: dict[str, TableAsset] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def iter_block_items(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", child
        elif child.tag == qn("w:tbl"):
            yield "table", child


def paragraph_from_element(doc: Document, element):
    from docx.text.paragraph import Paragraph
    return Paragraph(element, doc)


def table_from_element(doc: Document, element):
    from docx.table import Table as DocxTable
    return DocxTable(element, doc)


def safe_filename(name: str) -> str:
    name = re.sub(r"[^\w\-.]+", "_", name, flags=re.UNICODE).strip("_")
    return name or "file"


def extract_images_from_paragraph(
    doc: Document, paragraph, figures_dir: Path, counter_start: int
) -> tuple[list[FigureAsset], int]:
    figures = []
    counter = counter_start
    blips = paragraph._element.xpath(".//a:blip")
    for blip in blips:
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid or rid not in doc.part.related_parts:
            continue
        image_part = doc.part.related_parts[rid]
        original_name = Path(
            getattr(image_part, "partname", f"image_{counter}.png").basename
        ).name
        suffix = Path(original_name).suffix or ".png"
        figure_id = f"figure_{counter:03d}"
        filename = f"{figure_id}{suffix.lower()}"
        out_path = figures_dir / filename
        out_path.write_bytes(image_part.blob)
        figures.append(
            FigureAsset(figure_id=figure_id, filename=filename, source_path=out_path)
        )
        counter += 1
    return figures, counter


def table_rows(table) -> list[list[str]]:
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
    return rows


def guess_title(blocks: list[Block], fallback: str) -> str:
    for block in blocks:
        if block.type == "paragraph" and block.text.strip():
            if block.style in {"Title", "Heading 1"} or len(block.text.strip()) < 140:
                return block.text.strip()
    return fallback


def extract_docx(input_path: Path, work_dir: Path) -> ExtractedDocument:
    figures_dir = work_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(input_path)
    blocks: list[Block] = []
    figures: dict[str, FigureAsset] = {}
    tables: dict[str, TableAsset] = {}
    figure_counter = 1
    table_counter = 1

    for kind, element in iter_block_items(doc):
        if kind == "paragraph":
            para = paragraph_from_element(doc, element)
            text = para.text.strip()
            style = para.style.name if para.style is not None else "Normal"
            para_figures, figure_counter = extract_images_from_paragraph(
                doc, para, figures_dir, figure_counter
            )
            if text:
                blocks.append(Block(type="paragraph", text=text, style=style))
            for fig in para_figures:
                figures[fig.figure_id] = fig
                blocks.append(Block(type="figure", ref_id=fig.figure_id))
        elif kind == "table":
            table = table_from_element(doc, element)
            table_id = f"table_{table_counter:03d}"
            table_counter += 1
            tables[table_id] = TableAsset(table_id=table_id, rows=table_rows(table))
            blocks.append(Block(type="table", ref_id=table_id))

    add_captions_from_neighbor_paragraphs(blocks, figures, tables)
    title = guess_title(blocks, input_path.stem)
    return ExtractedDocument(title=title, blocks=blocks, figures=figures, tables=tables)


def is_caption_text(text: str, prefix: str) -> bool:
    if prefix == "figure":
        prefix_re = r"^(figure|fig\.|figura)\s*\d*[:\.-]?\s+"
    elif prefix == "table":
        prefix_re = r"^(table|tabela)\s*\d*[:\.-]?\s+"
    else:
        prefix_re = r"^(figure|fig\.|figura|tabela|table)\s*\d*[:\.-]?\s+"
    return bool(re.match(prefix_re, text.strip(), flags=re.IGNORECASE))


def add_captions_from_neighbor_paragraphs(
    blocks: list[Block],
    figures: dict[str, FigureAsset],
    tables: dict[str, TableAsset],
) -> None:
    for i, block in enumerate(blocks):
        if block.type not in {"figure", "table"}:
            continue
        target = (
            figures.get(block.ref_id)
            if block.type == "figure"
            else tables.get(block.ref_id)
        )
        if target is None:
            continue
        caption_type = "figure" if block.type == "figure" else "table"
        candidates = []
        if i + 1 < len(blocks) and blocks[i + 1].type == "paragraph":
            candidates.append(blocks[i + 1].text)
        if i - 1 >= 0 and blocks[i - 1].type == "paragraph":
            candidates.append(blocks[i - 1].text)
        for candidate in candidates:
            if is_caption_text(candidate, caption_type):
                target.caption = candidate
                break


# ---------------------------------------------------------------------------
# Markup for revision
# ---------------------------------------------------------------------------

def document_to_revision_text(extracted: ExtractedDocument) -> str:
    lines = []
    for block in extracted.blocks:
        if block.type == "paragraph":
            lines.append(block.text)
        elif block.type == "table":
            lines.append(f"[[TABLE:{block.ref_id}]]")
        elif block.type == "figure":
            lines.append(f"[[FIGURE:{block.ref_id}]]")
    return "\n\n".join(lines).strip()


def split_text(text: str, max_chars: int) -> list[str]:
    paragraphs = re.split(r"\n\s*\n", text)
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        candidate = f"{current}\n\n{paragraph}".strip()
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                chunks.append(current.strip())
            current = paragraph
    if current:
        chunks.append(current.strip())
    return chunks


# FIX: call_claude now accepts an optional system_prompt parameter.
# Previously hardcoded SYSTEM_PROMPT, which meant VISUAL_AUDIT_PROMPT
# was being passed as a user message instead of the system role.
def call_claude(prompt: str, attempts: int = 3, system: str | None = None) -> str:
    if anthropic is None:
        raise RuntimeError(
            "Missing package: anthropic. Run: pip install -r requirements.txt"
        )
    if not os.getenv("ANTHROPIC_API_KEY"):
        raise RuntimeError("Missing ANTHROPIC_API_KEY in .env")

    client = anthropic.Anthropic()
    system_prompt = system if system is not None else SYSTEM_PROMPT
    last_error = None

    for i in range(attempts):
        try:
            response = client.messages.create(
                model=CONFIG["model"],
                max_tokens=CONFIG["max_output_tokens"],
                temperature=CONFIG["temperature"],
                system=system_prompt,
                messages=[{"role": "user", "content": prompt}],
            )
            return "\n".join(
                block.text
                for block in response.content
                if getattr(block, "type", None) == "text"
            ).strip()
        except Exception as exc:
            last_error = exc
            if i < attempts - 1:
                time.sleep(3 * (i + 1))

    raise RuntimeError(
        f"Claude call failed after {attempts} attempts: {last_error}"
    )


def revise_chunk(chunk: str, language: str, part: int, total: int) -> str:
    prompt = f"""
Target language: {LANGUAGE_NAME[language]}

Document section: {part} of {total}

Revise this section as polished academic writing for a lay but informed audience.

Critical placeholder rule:
Keep every placeholder exactly unchanged, including brackets and IDs. Examples:
  [[TABLE:table_001]]
  [[FIGURE:figure_001]]

Original section:

{chunk}
"""
    return call_claude(prompt)


# FIX: harmonize now chunks large documents to avoid context-window overflow.
# The previous version sent the entire document as a single API call, which
# silently truncates or fails on documents longer than ~30 pages.
def harmonize(text: str, language: str) -> str:
    chunks = split_text(text, CONFIG["max_chunk_chars"])

    if len(chunks) == 1:
        prompt = f"""
Target language: {LANGUAGE_NAME[language]}

Do a final editorial harmonization of this document.

Tasks:
- Make terminology consistent throughout.
- Smooth transitions between paragraphs and sections.
- Remove duplicated ideas.
- Preserve all facts, numbers, and [[TABLE:...]] / [[FIGURE:...]] placeholders exactly.
- Return only the final document text.

Text:

{chunks[0]}
"""
        return call_claude(prompt)

    harmonized_chunks = []
    for i, chunk in enumerate(chunks):
        context_before = chunks[i - 1][-600:] if i > 0 else ""
        context_after = chunks[i + 1][:600] if i < len(chunks) - 1 else ""

        prompt = f"""
Target language: {LANGUAGE_NAME[language]}

Harmonize this document section (section {i + 1} of {len(chunks)}).

Tasks:
- Make terminology consistent with the surrounding sections.
- Smooth the opening and closing transitions.
- Remove ideas that duplicate adjacent sections.
- Preserve all facts, numbers, and [[TABLE:...]] / [[FIGURE:...]] placeholders exactly.
- Return only the harmonized section text.

Context from previous section (last 600 characters):
{context_before}

Current section to harmonize:
{chunk}

Context from next section (first 600 characters):
{context_after}
"""
        harmonized_chunks.append(call_claude(prompt))

    return "\n\n".join(harmonized_chunks)


def revise_document_text(marked_text: str, language: str) -> str:
    chunks = split_text(marked_text, CONFIG["max_chunk_chars"])
    revised = []
    for i, chunk in enumerate(chunks, start=1):
        print(f"  Revising {language}: section {i}/{len(chunks)}")
        revised.append(revise_chunk(chunk, language, i, len(chunks)))
    print(f"  Harmonizing {language} final version")
    return harmonize("\n\n".join(revised), language)


def make_manual_prompts(marked_text: str, out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    for language in ["pt-BR", "en-US"]:
        chunks = split_text(marked_text, CONFIG["max_chunk_chars"])
        for i, chunk in enumerate(chunks, start=1):
            prompt = f"""{SYSTEM_PROMPT}

Target language: {LANGUAGE_NAME[language]}

Document section: {i} of {len(chunks)}

Revise this section. Keep all [[TABLE:...]] and [[FIGURE:...]] placeholders exactly unchanged.

Original section:

{chunk}
"""
            (out_dir / f"{language}_section_{i:03d}.txt").write_text(
                prompt, encoding="utf-8"
            )


# ---------------------------------------------------------------------------
# DOCX table borders -- three-line academic style
# ---------------------------------------------------------------------------

# FIX: replaced "Table Grid" (all borders including verticals) with a
# three-line academic style: top rule, below-header rule, bottom rule only.
# "Table Grid" contradicts the three-line table standard recommended in the
# VISUAL_AUDIT_PROMPT produced by this same tool.

def _set_cell_border(cell, **kwargs):
    """Set borders on a single cell. kwargs: top, bottom, left, right = (val, sz, color)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, attrs in kwargs.items():
        tag = tcBorders.find(qn(f"w:{side}"))
        if tag is None:
            tag = OxmlElement(f"w:{side}")
            tcBorders.append(tag)
        val, sz, color = attrs
        tag.set(qn("w:val"), val)
        tag.set(qn("w:sz"), str(sz))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def _apply_three_line_borders(table, n_rows: int) -> None:
    """
    Apply three-line table borders:
      - top rule (1.5 pt) on all cells in row 0
      - mid rule (0.75 pt) below header (all cells in row 0)
      - bottom rule (1.5 pt) on all cells in last row
      - no vertical or internal horizontal borders anywhere
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            borders = {}
            if i == 0:
                borders["top"] = ("single", 12, "000000")     # 1.5 pt top rule
                borders["bottom"] = ("single", 6, "000000")   # 0.75 pt mid rule
            elif i == n_rows - 1:
                borders["bottom"] = ("single", 12, "000000")  # 1.5 pt bottom rule
            # Suppress all other borders explicitly
            for side in ("left", "right", "insideH", "insideV"):
                borders[side] = ("none", 0, "auto")
            if i != 0:
                borders["top"] = ("none", 0, "auto")
                borders["bottom"] = ("none", 0, "auto")
            _set_cell_border(cell, **borders)
            # Re-apply bottom for last row
            if i == n_rows - 1:
                _set_cell_border(cell, bottom=("single", 12, "000000"))


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

def add_table_to_docx(
    doc: Document, table_asset: TableAsset, caption: str | None = None
) -> None:
    caption_text = caption or table_asset.caption or f"Table {table_asset.table_id[-3:]}"
    p = doc.add_paragraph()
    r = p.add_run(caption_text)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)

    rows = table_asset.rows or [[""]]
    n_cols = max(len(row) for row in rows)
    n_rows = len(rows)

    tbl = doc.add_table(rows=0, cols=n_cols)
    # Use Plain Table style as base (no pre-existing borders)
    tbl.style = "Table Normal"

    for i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            cells[j].text = cell_text
            for run in cells[j].paragraphs[0].runs:
                run.font.name = "Arial" if i == 0 else "Palatino Linotype"
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

    _apply_three_line_borders(tbl, n_rows)
    doc.add_paragraph()


def add_figure_to_docx(
    doc: Document, figure: FigureAsset, caption: str | None = None
) -> None:
    if figure.source_path.exists():
        try:
            doc.add_picture(str(figure.source_path), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Could not insert figure: {figure.filename}]")
    caption_text = caption or figure.caption or f"Figure {figure.figure_id[-3:]}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption_text)
    r.italic = True
    r.font.size = Pt(9)
    doc.add_paragraph()


# FIX: is_heading_line now uses much stricter heuristics to avoid
# misclassifying short paragraphs as section headings.
# Max length reduced from 120 to 80 characters.
# Additional checks: no multiple sentences, no lowercase start, no list items.
def is_heading_line(line: str, style: str = "Normal") -> bool:
    """True if the line should be formatted as a document heading."""
    # If we have the original Word style, use it as ground truth
    if style in HEADING_STYLES:
        return True

    clean = line.strip()
    if not clean:
        return False
    if clean.startswith("[["):
        return False

    # Length guard: real headings are short
    if len(clean) > 80:
        return False

    # Headings do not end with sentence-closing punctuation
    if clean.endswith((".", ":", ",", ";")):
        return False

    # Headings do not contain multiple sentences
    if re.search(r"\.\s+[A-Z]", clean):
        return False

    # Headings do not start with a lowercase letter
    if clean[0].islower():
        return False

    # Numbered list items are not headings
    if re.match(r"^\d+[\.\)]\s", clean):
        return False

    # Parenthetical-heavy lines are body text
    if clean.count("(") >= 2:
        return False

    return True


def write_docx_from_revised(
    text: str, extracted: ExtractedDocument, output_path: Path, title: str
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Palatino Linotype"
    normal.font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(26, 107, 58)

    for block in re.split(r"\n\s*\n", text.strip()):
        block = block.strip()
        if not block:
            continue

        table_match = re.fullmatch(r"\[\[TABLE:(table_\d{3})\]\]", block)
        fig_match = re.fullmatch(r"\[\[FIGURE:(figure_\d{3})\]\]", block)

        if table_match:
            table_id = table_match.group(1)
            if table_id in extracted.tables:
                add_table_to_docx(doc, extracted.tables[table_id])
            continue

        if fig_match:
            figure_id = fig_match.group(1)
            if figure_id in extracted.figures:
                add_figure_to_docx(doc, extracted.figures[figure_id])
            continue

        if is_heading_line(block):
            h = doc.add_heading(block, level=1)
            for run in h.runs:
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(26, 107, 58)
        else:
            p = doc.add_paragraph(block)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.12

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# LaTeX and PDF output
# ---------------------------------------------------------------------------

# FIX: latex_escape now uses a placeholder to prevent the backslash-to-
# \textbackslash{} substitution from having its curly braces re-escaped
# by the subsequent { -> \{ and } -> \} substitutions.
def latex_escape(text: str) -> str:
    if text is None:
        return ""
    text = str(text)
    _BS = "\x00BS\x00"  # internal placeholder for backslash
    text = text.replace("\\", _BS)
    text = text.replace("&", r"\&")
    text = text.replace("%", r"\%")
    text = text.replace("$", r"\$")
    text = text.replace("#", r"\#")
    text = text.replace("_", r"\_")
    text = text.replace("{", r"\{")
    text = text.replace("}", r"\}")
    text = text.replace("~", r"\textasciitilde{}")
    text = text.replace("^", r"\textasciicircum{}")
    # Restore backslash as \textbackslash{} after all other substitutions
    text = text.replace(_BS, r"\textbackslash{}")
    # Unicode math symbols
    text = text.replace("\u2265", r"$\geq$")
    text = text.replace("\u2264", r"$\leq$")
    text = text.replace("\u2248", r"$\approx$")
    text = text.replace("\u00d7", r"$\times$")
    return text


# FIX: table_to_latex now estimates per-column widths based on actual content
# length instead of assigning uniform width to every column.
# A table with one long text column and three short numeric columns
# will now render correctly instead of with identical narrow columns.
def _estimate_col_widths(rows: list[list[str]], n_cols: int) -> list[float]:
    """Return fractional textwidth for each column based on max cell length."""
    max_lens: list[int] = []
    for j in range(n_cols):
        col_max = max(
            (len(row[j]) if j < len(row) else 0) for row in rows
        )
        max_lens.append(max(col_max, 4))  # floor at 4 chars to avoid zero-width
    total = sum(max_lens)
    available = 0.90  # fraction of textwidth reserved for table
    return [available * ln / total for ln in max_lens]


def table_to_latex(asset: TableAsset) -> str:
    rows = asset.rows or [[""]]
    n_cols = max(len(row) for row in rows)
    widths = _estimate_col_widths(rows, n_cols)
    colspec = " ".join(f"L{{{w:.2f}\\textwidth}}" for w in widths)
    caption = latex_escape(asset.caption or f"Table {asset.table_id[-3:]}")
    label = asset.table_id

    lines = [
        r"\begin{table}[H]",
        r"\centering",
        rf"\caption{{{caption}\label{{tab:{label}}}}}",
        rf"\begin{{tabularx}}{{\textwidth}}{{{colspec}}}",
        r"\toprule",
    ]
    for i, row in enumerate(rows):
        cells = [latex_escape(row[j] if j < len(row) else "") for j in range(n_cols)]
        if i == 0:
            cells = [rf"\textbf{{{c}}}" for c in cells]
        lines.append(" & ".join(cells) + r"\\")
        if i == 0:
            lines.append(r"\midrule")
    lines.extend([r"\bottomrule", r"\end{tabularx}", r"\end{table}"])
    return "\n".join(lines)


def figure_to_latex(asset: FigureAsset, figures_rel_dir: str = "figures") -> str:
    caption = latex_escape(asset.caption or f"Figure {asset.figure_id[-3:]}")
    label = asset.figure_id
    return "\n".join([
        r"\begin{figure}[H]",
        r"\centering",
        rf"\includegraphics[width=0.88\linewidth]{{{figures_rel_dir}/{asset.filename}}}",
        rf"\caption{{{caption}\label{{fig:{label}}}}}",
        r"\end{figure}",
    ])


# FIX: text_to_latex_body now detects numbered headings (e.g. "1 Introduction",
# "2.1 Methods") to assign correct hierarchy (\section, \subsection,
# \subsubsection) instead of treating only the first heading as \section
# and all subsequent ones as \subsection regardless of depth.
def _latex_heading_level(line: str) -> int:
    """
    Detect heading level from numbering prefix.
    '1 Introduction' -> 1
    '1.1 Methods'    -> 2
    '1.1.1 Sub'      -> 3
    Unnumbered       -> 1
    """
    m = re.match(r"^(\d+(?:\.\d+)*)\s", line.strip())
    if m:
        depth = m.group(1).count(".") + 1
        return min(depth, 3)
    return 1


def text_to_latex_body(text: str, extracted: ExtractedDocument) -> str:
    lines = []
    level_cmds = {1: "section", 2: "subsection", 3: "subsubsection"}

    for raw_block in re.split(r"\n\s*\n", text.strip()):
        block = raw_block.strip()
        if not block:
            continue

        table_match = re.fullmatch(r"\[\[TABLE:(table_\d{3})\]\]", block)
        fig_match = re.fullmatch(r"\[\[FIGURE:(figure_\d{3})\]\]", block)

        if table_match:
            table_id = table_match.group(1)
            if table_id in extracted.tables:
                lines.append(table_to_latex(extracted.tables[table_id]))
            continue

        if fig_match:
            figure_id = fig_match.group(1)
            if figure_id in extracted.figures:
                lines.append(figure_to_latex(extracted.figures[figure_id]))
            continue

        if is_heading_line(block):
            level = _latex_heading_level(block)
            cmd = level_cmds.get(level, "section")
            lines.append(rf"\{cmd}{{{latex_escape(block)}}}")
        else:
            lines.append(latex_escape(block))

    return "\n\n".join(lines)


def default_latex_template() -> str:
    return r"""
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{mathpazo}
\usepackage[scaled=0.92]{helvet}
\usepackage[a4paper,margin=1in]{geometry}
\usepackage{xcolor}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{tabularx}
\usepackage{array}
\usepackage{float}
\usepackage{caption}
\usepackage{titlesec}
\usepackage{fancyhdr}
\usepackage{setspace}
\usepackage{microtype}
\usepackage{hyperref}

\definecolor{mgreen}{RGB}{26,107,58}
\definecolor{mblue}{RGB}{0,103,172}

\newcolumntype{L}[1]{>{\raggedright\arraybackslash}p{#1}}
\newcolumntype{Y}{>{\raggedright\arraybackslash}X}

\titleformat{\section}{\Large\bfseries\sffamily\color{mgreen}}{}{0pt}{}
\titleformat{\subsection}{\large\bfseries\sffamily\color{mblue}}{}{0pt}{}
\titleformat{\subsubsection}{\normalsize\bfseries\sffamily}{}{0pt}{}

\captionsetup{font=small,labelfont=bf}
\hypersetup{colorlinks=true,linkcolor=mblue,urlcolor=mblue,citecolor=mblue}

\pagestyle{fancy}
\fancyhf{}
\lhead{AI Word Editor}
\rhead{\thepage}

\setstretch{1.08}
\setlength{\parskip}{0.55em}
\setlength{\parindent}{0pt}
""".strip()


def build_latex_document(text: str, extracted: ExtractedDocument, title: str) -> str:
    preamble = default_latex_template()
    body = text_to_latex_body(text, extracted)
    return "\n".join([
        r"\documentclass[11pt,a4paper]{article}",
        preamble,
        r"\begin{document}",
        rf"\title{{{latex_escape(title)}}}",
        r"\author{}",
        r"\date{}",
        r"\maketitle",
        body,
        r"\end{document}",
        "",
    ])


def compile_latex_pdf(tex_path: Path, pdf_out: Path) -> bool:
    if shutil.which("pdflatex") is None:
        return False
    for pass_n in [1, 2]:
        print(f"  Compiling LaTeX PDF pass {pass_n}/2: {tex_path.name}")
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=tex_path.parent,
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            log = tex_path.with_suffix(".log")
            print(f"  LaTeX failed. See: {log}")
            return False
    built_pdf = tex_path.with_suffix(".pdf")
    if built_pdf.exists():
        pdf_out.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(built_pdf, pdf_out)
        return True
    return False


# FIX: write_reportlab_pdf now uses LINEABOVE / LINEBELOW for a three-line
# academic table style instead of GRID (which draws all borders including
# verticals). This aligns the DOCX and PDF outputs with the three-line
# standard recommended in the tool's own VISUAL_AUDIT_PROMPT.
def write_reportlab_pdf(
    text: str, extracted: ExtractedDocument, output_path: Path, title: str
) -> None:
    if colors is None:
        raise RuntimeError(
            "ReportLab is not installed and pdflatex is unavailable. "
            "Run: pip install reportlab"
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="DocTitle", parent=styles["Title"],
        fontName="Helvetica-Bold", fontSize=16,
        textColor=colors.HexColor("#1A6B3A"), spaceAfter=16,
    ))
    styles.add(ParagraphStyle(
        name="HeadingGreen", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=13,
        textColor=colors.HexColor("#1A6B3A"), spaceBefore=12, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="HeadingBlue", parent=styles["Heading2"],
        fontName="Helvetica-Bold", fontSize=11,
        textColor=colors.HexColor("#006780"), spaceBefore=8, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="BodyReadable", parent=styles["BodyText"],
        fontName="Times-Roman", fontSize=10.5, leading=14, spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="Caption", parent=styles["BodyText"],
        fontName="Times-Italic", fontSize=8.5, alignment=1,
        spaceBefore=4, spaceAfter=10,
    ))

    story = [Paragraph(html.escape(title), styles["DocTitle"])]

    for raw_block in re.split(r"\n\s*\n", text.strip()):
        block = raw_block.strip()
        if not block:
            continue

        table_match = re.fullmatch(r"\[\[TABLE:(table_\d{3})\]\]", block)
        fig_match = re.fullmatch(r"\[\[FIGURE:(figure_\d{3})\]\]", block)

        if table_match:
            asset = extracted.tables.get(table_match.group(1))
            if asset:
                rows = asset.rows or [[""]]
                tbl = Table(rows, repeatRows=1)
                tbl.setStyle(TableStyle([
                    # Header row background
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEEEEE")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    # Three-line style: top rule, mid rule, bottom rule
                    ("LINEABOVE", (0, 0), (-1, 0), 1.2, colors.black),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.black),
                    ("LINEBELOW", (0, -1), (-1, -1), 1.2, colors.black),
                    # Padding
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ]))
                caption = Paragraph(
                    html.escape(asset.caption or f"Table {asset.table_id[-3:]}"),
                    styles["Caption"],
                )
                story.append(KeepTogether([caption, tbl, Spacer(1, 0.12 * inch)]))
            continue

        if fig_match:
            asset = extracted.figures.get(fig_match.group(1))
            if asset and asset.source_path.exists():
                try:
                    img = RLImage(str(asset.source_path))
                    max_width = 6.6 * inch
                    if img.drawWidth > max_width:
                        ratio = max_width / img.drawWidth
                        img.drawWidth *= ratio
                        img.drawHeight *= ratio
                    img.hAlign = "CENTER"
                    cap = Paragraph(
                        html.escape(asset.caption or f"Figure {asset.figure_id[-3:]}"),
                        styles["Caption"],
                    )
                    story.append(KeepTogether([img, cap]))
                except Exception:
                    story.append(Paragraph(
                        f"[Could not render figure: {html.escape(asset.filename)}]",
                        styles["BodyReadable"],
                    ))
            continue

        if is_heading_line(block):
            level = _latex_heading_level(block)
            style_name = "HeadingGreen" if level == 1 else "HeadingBlue"
            story.append(Paragraph(html.escape(block), styles[style_name]))
        else:
            story.append(Paragraph(html.escape(block), styles["BodyReadable"]))

    doc.build(story)


def write_pdf_outputs(
    text: str,
    extracted: ExtractedDocument,
    build_dir: Path,
    output_pdf: Path,
    title: str,
) -> None:
    tex_path = build_dir / f"{output_pdf.stem}.tex"
    tex_path.write_text(
        build_latex_document(text, extracted, title), encoding="utf-8"
    )
    ok = compile_latex_pdf(tex_path, output_pdf)
    if not ok:
        print("  pdflatex unavailable or failed. Building fallback ReportLab PDF.")
        write_reportlab_pdf(text, extracted, output_pdf, title)


# ---------------------------------------------------------------------------
# Visual communication audit
# ---------------------------------------------------------------------------

def surrounding_text_for_visual(
    blocks: list[Block], idx: int, window: int = 2
) -> str:
    start = max(0, idx - window)
    end = min(len(blocks), idx + window + 1)
    parts = []
    for j in range(start, end):
        b = blocks[j]
        if b.type == "paragraph" and b.text.strip():
            parts.append(b.text.strip())
        elif b.type == "table":
            parts.append(f"[[TABLE:{b.ref_id}]]")
        elif b.type == "figure":
            parts.append(f"[[FIGURE:{b.ref_id}]]")
    return "\n\n".join(parts)


def visual_inventory(extracted: ExtractedDocument, max_table_rows: int = 12) -> str:
    lines = [
        f"Document title: {extracted.title}",
        "",
        "Detected visuals and surrounding context:",
    ]
    visual_n = 1
    for i, block in enumerate(extracted.blocks):
        if block.type == "table" and block.ref_id in extracted.tables:
            asset = extracted.tables[block.ref_id]
            rows = asset.rows or []
            n_cols = max((len(r) for r in rows), default=0)
            lines += [
                "",
                f"VISUAL {visual_n}: table -- {asset.caption or block.ref_id}",
                f"Rows: {len(rows)} | Columns: {n_cols}",
                f"Caption: {asset.caption or 'No caption detected'}",
                "Surrounding text:",
                surrounding_text_for_visual(extracted.blocks, i),
                "Table preview:",
            ]
            for row in rows[:max_table_rows]:
                lines.append(" | ".join(row))
            if len(rows) > max_table_rows:
                lines.append(
                    f"[Table truncated: {len(rows) - max_table_rows} additional rows]"
                )
            visual_n += 1
        elif block.type == "figure" and block.ref_id in extracted.figures:
            asset = extracted.figures[block.ref_id]
            lines += [
                "",
                f"VISUAL {visual_n}: figure/map/chart -- {asset.caption or block.ref_id}",
                f"File: {asset.filename}",
                f"Caption: {asset.caption or 'No caption detected'}",
                "Surrounding text:",
                surrounding_text_for_visual(extracted.blocks, i),
            ]
            visual_n += 1

    if visual_n == 1:
        lines.append("No embedded figures or tables were detected in the Word file.")
    return "\n".join(lines)


def write_markdown_as_docx(
    markdown_text: str, output_path: Path, title: str
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Palatino Linotype"
    normal.font.size = Pt(10.5)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(26, 107, 58)

    AUDIT_SECTION_RE = re.compile(
        r"^(VISUAL \d+:|DIAGNOSIS|CHART/MAP TYPE|DATA INTEGRITY|"
        r"TITLE AND LABELS|COLOR|ANNOTATIONS|CAPTION|MAP-SPECIFIC|"
        r"TABLE-SPECIFIC|CROSS-REFERENCE|PRIORITY|VISUAL COHERENCE SUMMARY)"
    )

    for block in re.split(r"\n\s*\n", markdown_text.strip()):
        clean = block.strip()
        if not clean:
            continue
        if clean.startswith("# "):
            h = doc.add_heading(clean[2:].strip(), level=1)
            for run in h.runs:
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(26, 107, 58)
        elif clean.startswith("## "):
            h = doc.add_heading(clean[3:].strip(), level=2)
            for run in h.runs:
                run.font.name = "Arial"
        elif AUDIT_SECTION_RE.match(clean):
            p = doc.add_paragraph()
            r = p.add_run(clean)
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph(clean)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.line_spacing = 1.10

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def write_markdown_pdf(
    markdown_text: str, output_path: Path, title: str
) -> None:
    if colors is None:
        return
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="AuditTitle", parent=styles["Title"],
        fontName="Helvetica-Bold", fontSize=16,
        textColor=colors.HexColor("#1A6B3A"), spaceAfter=16,
    ))
    styles.add(ParagraphStyle(
        name="AuditHeading", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=12,
        textColor=colors.HexColor("#1A6B3A"), spaceBefore=10, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="AuditBody", parent=styles["BodyText"],
        fontName="Times-Roman", fontSize=9.8, leading=13, spaceAfter=6,
    ))

    AUDIT_SECTION_RE = re.compile(
        r"^(VISUAL \d+:|DIAGNOSIS|CHART/MAP TYPE|DATA INTEGRITY|"
        r"TITLE AND LABELS|COLOR|ANNOTATIONS|CAPTION|MAP-SPECIFIC|"
        r"TABLE-SPECIFIC|CROSS-REFERENCE|PRIORITY|VISUAL COHERENCE SUMMARY)"
    )

    story = [Paragraph(html.escape(title), styles["AuditTitle"])]
    for block in re.split(r"\n\s*\n", markdown_text.strip()):
        clean = block.strip()
        if not clean:
            continue
        if (
            clean.startswith("# ")
            or clean.startswith("## ")
            or AUDIT_SECTION_RE.match(clean)
        ):
            story.append(Paragraph(
                html.escape(clean.lstrip("# ").strip()),
                styles["AuditHeading"],
            ))
        else:
            story.append(Paragraph(
                html.escape(clean).replace("\n", "<br/>"),
                styles["AuditBody"],
            ))
    doc.build(story)


# FIX: build_visual_audit now passes VISUAL_AUDIT_PROMPT as the system
# parameter to call_claude, not as part of the user message. Previously
# the audit instructions were concatenated into the user prompt while
# SYSTEM_PROMPT (the text editing instructions) remained in the system role,
# causing the model to apply text-editing rules to a visual audit task.
def build_visual_audit(
    extracted: ExtractedDocument, language: str, full_text: str
) -> str:
    if not extracted.tables and not extracted.figures:
        return "No embedded figures or tables were detected in the Word file."

    inventory = visual_inventory(extracted)
    # Truncate full_text to avoid exceeding context window
    text_preview = full_text[:8000] + ("..." if len(full_text) > 8000 else "")

    prompt = f"""
Language for audit report: {LANGUAGE_NAME.get(language, language)}

{inventory}

Full revised document text for cross-reference (first 8000 characters):

{text_preview}
"""
    return call_claude(prompt, system=VISUAL_AUDIT_PROMPT)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "AI Word Editor: revise a .docx file for a mixed academic and "
            "lay audience. Outputs revised DOCX and PDF in Brazilian Portuguese "
            "and American English."
        )
    )
    parser.add_argument("input", help="Path to input .docx file")
    parser.add_argument(
        "--manual",
        action="store_true",
        help="Generate copy-paste prompts instead of calling the API",
    )
    parser.add_argument(
        "--audit",
        action="store_true",
        help="Run visual communication audit after text revision",
    )
    parser.add_argument(
        "--audit-only",
        action="store_true",
        help="Run only the visual audit, skip text revision",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    if input_path.suffix.lower() != ".docx":
        print(f"Error: expected a .docx file, got: {input_path.suffix}", file=sys.stderr)
        sys.exit(1)

    stem = safe_filename(input_path.stem)
    out_dir = CONFIG["output_dir"] / stem
    work_dir = out_dir / "_work"
    work_dir.mkdir(parents=True, exist_ok=True)

    print(f"Extracting: {input_path.name}")
    extracted = extract_docx(input_path, work_dir)
    print(f"Title    : {extracted.title}")
    print(
        f"Blocks   : {len(extracted.blocks)} paragraphs/placeholders | "
        f"Tables: {len(extracted.tables)} | Figures: {len(extracted.figures)}"
    )

    marked_text = document_to_revision_text(extracted)

    if args.manual:
        prompts_dir = out_dir / f"{stem}_manual_prompts"
        make_manual_prompts(marked_text, prompts_dir)
        print(f"\nManual prompts written to: {prompts_dir}")
        return

    revised_texts: dict[str, str] = {}

    if not args.audit_only:
        for language in ["pt-BR", "en-US"]:
            print(f"\nRevising ({language})")
            revised = revise_document_text(marked_text, language)
            revised_texts[language] = revised

            lang_code = language.replace("-", "_").lower()
            docx_out = out_dir / f"{stem}_revised_{lang_code}.docx"
            build_dir = work_dir / lang_code
            build_dir.mkdir(parents=True, exist_ok=True)

            # Copy figures into the LaTeX build directory
            figs_build = build_dir / "figures"
            figs_build.mkdir(exist_ok=True)
            for fig in extracted.figures.values():
                if fig.source_path.exists():
                    shutil.copy2(fig.source_path, figs_build / fig.filename)

            print(f"  Writing DOCX: {docx_out.name}")
            write_docx_from_revised(revised, extracted, docx_out, extracted.title)

            pdf_out = out_dir / f"{stem}_revised_{lang_code}.pdf"
            print(f"  Writing PDF : {pdf_out.name}")
            write_pdf_outputs(revised, extracted, build_dir, pdf_out, extracted.title)

        print(f"\nOutputs written to: {out_dir}")

    if args.audit or args.audit_only:
        audit_language = "pt-BR"
        audit_text = revised_texts.get(audit_language, marked_text)

        print(f"\nRunning visual audit ({audit_language})")
        audit_result = build_visual_audit(extracted, audit_language, audit_text)

        audit_docx = out_dir / f"{stem}_visual_audit.docx"
        audit_pdf = out_dir / f"{stem}_visual_audit.pdf"

        print(f"  Writing audit DOCX: {audit_docx.name}")
        write_markdown_as_docx(
            audit_result, audit_docx, f"Visual Audit: {extracted.title}"
        )
        print(f"  Writing audit PDF : {audit_pdf.name}")
        write_markdown_pdf(
            audit_result, audit_pdf, f"Visual Audit: {extracted.title}"
        )
        print(f"\nAudit written to: {out_dir}")


if __name__ == "__main__":
    main()
